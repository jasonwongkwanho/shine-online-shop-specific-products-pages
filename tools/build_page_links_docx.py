from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "page-links.docx"

BASE_URL = "https://jasonwongkwanho.github.io/shine-online-shop-specific-products-pages/"
ROWS = [
    ("眼鏡布_名畫喵系列", "products/masterpiece-cat-cloth.html", f"{BASE_URL}products/masterpiece-cat-cloth.html"),
    ("眼鏡布_犬雕像系列", "products/dog-statue-cloth.html", f"{BASE_URL}products/dog-statue-cloth.html"),
    ("立體咭_環球系列", "products/pop-up-world.html", f"{BASE_URL}products/pop-up-world.html"),
    ("立體咭_節慶系列", "products/pop-up-festival.html", f"{BASE_URL}products/pop-up-festival.html"),
    ("立體咭_港式情懷", "products/pop-up-hong-kong.html", f"{BASE_URL}products/pop-up-hong-kong.html"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F766E")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_run(run, bold=False, color=None, size=11):
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("網尚店產品展示頁連結表")
    style_run(title_run, bold=True, color="0B2545", size=20)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run("以下完整 Link 可直接貼入 Canva Website。")
    style_run(run, color="334155", size=10.5)

    base = doc.add_paragraph()
    base.paragraph_format.space_after = Pt(12)
    run = base.add_run(f"GitHub Pages 根網址：{BASE_URL}")
    style_run(run, color="65758B", size=10)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["產品系列", "頁面檔案", "相對 Link", "Canva 可用完整 Link"]
    widths = [2200, 2440, 2200, 2520]

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        style_run(run, bold=True, color="1F4D78", size=10)

    for product, relative, full in ROWS:
        cells = table.add_row().cells
        values = [product, relative, relative, full]
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx == 3:
                add_hyperlink(p, full, value)
            else:
                run = p.add_run(value)
                style_run(run, color="1F2933", size=9.5)

    set_table_geometry(table, widths)
    set_cell_margins(table)

    warning = doc.add_paragraph()
    warning.paragraph_format.space_before = Pt(10)
    warning.paragraph_format.space_after = Pt(0)
    run = warning.add_run("備註：頁面內圖片使用 Google Drive 圖片縮圖連結。正式公開前，請確保原 Drive 圖片或上層 folder 已設定為「知道連結的任何人均可查看」。")
    style_run(run, color="65758B", size=9.5)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
